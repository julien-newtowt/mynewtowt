"""Générateurs de documents Word (.docx) — backlog « DOCX generators ».

Regroupe les générateurs Word de la plateforme, en miroir de
``services.pdf_generator`` (qui produit les PDF via WeasyPrint) :

- ``build_offer_docx``           : offre commerciale (depuis ``RateOffer``).
- ``build_booking_note_docx``    : booking note contractuelle (depuis un
  ``BookingNote``, trame de type BIMCO CONLINEBOOKING, conditions générales au
  verso).
- ``build_bill_of_lading_docx``  : Bill of Lading / connaissement (depuis un
  ``Booking`` confirmé).

``python-docx`` est importé paresseusement : la dépendance n'est pas toujours
présente en dev et reste lourde à charger.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from datetime import UTC, datetime
from typing import Any

# Teal NEWTOWT (charte « Nouvelle Étoile ») — couleur d'accent des titres.
_TEAL = (0x0D, 0x59, 0x66)
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@dataclass
class DocxBytes:
    """Document Word sérialisé prêt à servir (téléchargement)."""

    docx: bytes
    filename: str
    mime: str = DOCX_MIME


# ---------------------------------------------------------------------------
# Helpers de mise en forme (chartés)
# ---------------------------------------------------------------------------


def _new_document():
    from docx import Document

    return Document()


def _teal_color():
    from docx.shared import RGBColor

    return RGBColor(*_TEAL)


def _title(doc, text: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    heading = doc.add_heading("", level=0)
    run = heading.add_run(text)
    run.font.color.rgb = _teal_color()
    run.font.size = Pt(20)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _section(doc, text: str):
    heading = doc.add_heading(text, level=2)
    if heading.runs:
        heading.runs[0].font.color.rgb = _teal_color()
    return heading


def _kv_table(doc, rows: list[tuple[str, str]]):
    """Table clé/valeur (2 colonnes, libellé en gras)."""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        if cells[0].paragraphs[0].runs:
            cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = value if value is not None else "—"
    return table


def _footer(doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc.add_paragraph()  # spacer
    para = doc.add_paragraph("NEWTOWT — Pioneer of wind-powered cargo since 2011 — www.newtowt.eu")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.color.rgb = _teal_color()
        run.font.size = Pt(9)
        run.italic = True


def _serialize(doc, filename: str) -> DocxBytes:
    buf = io.BytesIO()
    doc.save(buf)
    return DocxBytes(docx=buf.getvalue(), filename=filename)


def _fmt_eur(value) -> str:
    if value is None:
        return "—"
    return f"{value:,.2f} EUR".replace(",", " ")


def _fmt_date(value, fmt: str = "%d/%m/%Y") -> str:
    return value.strftime(fmt) if value else "—"


# ---------------------------------------------------------------------------
# Offre commerciale
# ---------------------------------------------------------------------------


def build_offer_docx(*, offer, client, leg) -> DocxBytes:
    """Offre commerciale Word depuis un ``RateOffer`` (+ client + leg optionnel)."""
    doc = _new_document()

    _title(doc, "OFFRE COMMERCIALE NEWTOWT")
    _ref_centered(doc, f"Référence : {offer.reference}")
    doc.add_paragraph()  # spacer

    _section(doc, "Client")
    # Contrat réel du modèle Client (commercial_clients) : ``name`` est la raison
    # sociale, le contact vit dans ``contact_name``/``contact_email``/``contact_phone``.
    client_rows = [("Société", client.name if client else "—")]
    if client and client.contact_name:
        client_rows.append(("Contact", client.contact_name))
    client_rows.append(("E-mail", client.contact_email if client else "—"))
    client_rows.append(("Téléphone", client.contact_phone if client else "—"))
    _kv_table(doc, client_rows)
    doc.add_paragraph()

    _section(doc, "Objet")
    doc.add_paragraph(offer.title or "—")
    doc.add_paragraph()

    _section(doc, "Itinéraire")
    if leg:
        doc.add_paragraph(
            f"Leg : {leg.leg_code}\n" f"ETD : {_fmt_date(leg.etd)}     ETA : {_fmt_date(leg.eta)}"
        )
    else:
        doc.add_paragraph("À confirmer")
    doc.add_paragraph()

    _section(doc, "Tarification")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for idx, label in enumerate(["Description", "Quantité", "Tarif unitaire", "Total"]):
        table.rows[0].cells[idx].text = label
        if table.rows[0].cells[idx].paragraphs[0].runs:
            table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
    rate = offer.proposed_rate_eur
    rate_str = f"{rate:,.2f} EUR/palette".replace(",", " ") if rate is not None else "—"
    cells = table.add_row().cells
    cells[0].text = "Fret palettes (voilier cargo)"
    cells[1].text = str(offer.estimated_palettes or 0)
    cells[2].text = rate_str
    cells[3].text = _fmt_eur(offer.total_eur)
    doc.add_paragraph()

    _section(doc, "Conditions")
    cond = doc.add_paragraph()
    cond.add_run(f"Validité : {_fmt_date(offer.valid_until)}\n")
    cond.add_run(
        "Ce prix inclut le transport par voilier cargo à propulsion vélique "
        "(zéro émission directe)."
    )

    if offer.notes:
        doc.add_paragraph()
        _section(doc, "Notes")
        doc.add_paragraph(offer.notes)

    _footer(doc)
    return _serialize(doc, f"Offre_{offer.reference}.docx")


def _ref_centered(doc, text: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


# ---------------------------------------------------------------------------
# Bill of Lading (connaissement)
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Dashboard environnemental — synthèse d'un voyage (LOT 12)
# ---------------------------------------------------------------------------

_PROP_LABELS_FR: dict[str, str] = {
    "velique_pur": "Vélique pur",
    "hybride": "Hybride (assistance vélique)",
    "mecanique": "Mécanique pur",
    "statique": "Statique / dérive",
}
_EVENT_LABELS_FR: dict[str, str] = {
    "departure": "Départ",
    "arrival": "Arrivée",
    "noon": "Noon",
    "anchoring_begin": "Début mouillage",
    "anchoring_end": "Fin mouillage",
}


def _fmt_num(value, dp: int = 2) -> str:
    if value is None:
        return "—"
    return f"{float(value):,.{dp}f}".replace(",", " ")


def _fmt_dt(value) -> str:
    if value is None:
        return "—"
    return value.strftime("%d/%m/%Y %H:%M")


def build_dashboard_voyage_docx(*, detail) -> DocxBytes:
    """Synthèse environnementale Word d'un voyage (``kpi_env.VoyageDetail``).

    Mêmes sections que l'export PDF : KPI conso/émissions, consommation vs
    cible, profil de propulsion (4 catégories + complétude), ROB (points de
    référence + chaîné + soutages), anomalies R14/R22. Calcul serveur
    exclusivement — ce générateur ne fait que restituer ``detail``."""
    ledger = detail.ledger
    doc = _new_document()

    _title(doc, "SYNTHÈSE ENVIRONNEMENTALE — VOYAGE")
    _ref_centered(doc, f"{detail.leg_code} · {detail.vessel_name or '—'}")
    doc.add_paragraph()

    _section(doc, "Voyage")
    route = "—"
    if detail.dep_port is not None and detail.arr_port is not None:
        route = f"{detail.dep_port.name} ({detail.dep_port.locode}) → {detail.arr_port.name} ({detail.arr_port.locode})"
    _kv_table(
        doc,
        [
            ("Navire", f"{detail.vessel_name or '—'} ({detail.vessel_code or '—'})"),
            ("Voyage · Leg", detail.leg_code),
            ("Route", route),
            ("Source des données", "événements" if detail.source == "events" else "noon legacy"),
            ("Durée (jours)", _fmt_num(detail.duration_days, 1)),
        ],
    )
    doc.add_paragraph()

    _section(doc, "Consommation & émissions")
    _kv_table(
        doc,
        [
            ("Consommation totale (t)", _fmt_num(ledger.conso_total_t, 3)),
            ("dont ME (t)", _fmt_num(ledger.conso_me_t, 3)),
            ("dont AE (t)", _fmt_num(ledger.conso_ae_t, 3)),
            ("CO₂ émis (t, TtW)", _fmt_num(ledger.co2_emitted_t, 3)),
            ("Distance (NM)", _fmt_num(ledger.distance_nm, 0)),
            ("Cargo B/L (t)", _fmt_num(ledger.cargo_bl_t, 1)),
            ("Cargo MRV (t)", _fmt_num(ledger.cargo_mrv_t, 1)),
            ("EF méthode A (gCO₂/t·km)", _fmt_num(ledger.ef_method_a, 2)),
            ("EF méthode B (gCO₂/t·km)", _fmt_num(ledger.ef_method_b, 2)),
            ("EF méthode C (gCO₂/t·km)", _fmt_num(ledger.ef_method_c, 2)),
        ],
    )
    doc.add_paragraph()

    _section(doc, "Consommation vs cible")
    conso = detail.conso
    if conso.daily_l_j is None:
        doc.add_paragraph(f"Consommation journalière : N/A — {conso.na_reason or ''}")
    else:
        verdict = "au-dessus" if conso.over_target else "en dessous ou égale"
        doc.add_paragraph(
            f"Consommation journalière : {_fmt_num(conso.daily_l_j, 0)} L/j "
            f"(cible {_fmt_num(conso.target_l_j, 0)} L/j — {verdict} du seuil, "
            f"écart {_fmt_num(conso.delta_pct, 1)} %)."
        )
    doc.add_paragraph()

    _section(doc, "Profil de propulsion (tranches de 4 h)")
    prop = detail.propulsion
    doc.add_paragraph(
        f"Complétude : {prop.filled_slots} tranches renseignées / "
        f"{prop.theoretical_slots} théoriques"
        + (
            f" ({_fmt_num(prop.completeness_pct, 1)} %)"
            if prop.completeness_pct is not None
            else ""
        )
        + ". Les tranches sans relevé sont exclues du dénominateur des pourcentages."
    )
    prop_table = doc.add_table(rows=1, cols=3)
    prop_table.style = "Table Grid"
    for idx, label in enumerate(["Catégorie", "Tranches", "% (des renseignées)"]):
        prop_table.rows[0].cells[idx].text = label
        if prop_table.rows[0].cells[idx].paragraphs[0].runs:
            prop_table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
    for seg in prop.segments:
        cells = prop_table.add_row().cells
        cells[0].text = _PROP_LABELS_FR.get(seg.category, seg.category)
        cells[1].text = str(seg.count)
        cells[2].text = f"{_fmt_num(seg.pct, 1)} %" if seg.pct is not None else "N/A"
    doc.add_paragraph()

    _section(doc, "ROB — points de référence & chaîné")
    rob_table = doc.add_table(rows=1, cols=4)
    rob_table.style = "Table Grid"
    for idx, label in enumerate(["Date (UTC)", "Événement", "ROB déclaré (t)", "ROB chaîné (t)"]):
        rob_table.rows[0].cells[idx].text = label
        if rob_table.rows[0].cells[idx].paragraphs[0].runs:
            rob_table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
    for p in detail.rob_chain:
        cells = rob_table.add_row().cells
        cells[0].text = _fmt_dt(p.datetime_utc)
        cells[1].text = _EVENT_LABELS_FR.get(p.event_type, p.event_type)
        cells[2].text = _fmt_num(p.rob_declared_t, 3)
        cells[3].text = _fmt_num(p.rob_calculated_t, 3)
    if detail.bunkers:
        doc.add_paragraph("Soutages (marqueurs) :")
        for b in detail.bunkers:
            doc.add_paragraph(
                f"  • BDN {b.bdn_number} — {_fmt_dt(b.delivery_datetime_utc)} — "
                f"{_fmt_num(b.mass_t, 3)} t @ {b.port_locode}",
                style="List Bullet",
            )
    doc.add_paragraph()

    _section(doc, "Anomalies qualité (R14 / R22)")
    if detail.quality:
        anom = doc.add_table(rows=1, cols=3)
        anom.style = "Table Grid"
        for idx, label in enumerate(["Règle", "Sévérité", "Message"]):
            anom.rows[0].cells[idx].text = label
            if anom.rows[0].cells[idx].paragraphs[0].runs:
                anom.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
        for q in detail.quality:
            cells = anom.add_row().cells
            cells[0].text = q.rule_id
            cells[1].text = q.severity_applied or "—"
            cells[2].text = (q.message or "—")[:300]
    else:
        doc.add_paragraph("Aucune anomalie R14/R22 ouverte sur ce voyage.")

    _footer(doc)
    return _serialize(doc, f"dashboard_voyage_{detail.leg_code}.docx")


def build_bill_of_lading_docx_from_pl(
    *,
    pl,
    batch,
    leg,
    vessel,
    pol,
    pod,
    bl_number,
    issued_at=None,
    shipped_on_board=None,
) -> DocxBytes:
    """Bill of Lading Word depuis un **lot de packing list** (rail registre).

    Remplace ``build_bill_of_lading_docx`` (rail booking), qui fabriquait un numéro
    à la volée sans jamais l'enregistrer.

    ⚠️ **Même règle d'honnêteté que le PDF** (cf. ``pdf_generator``) : tant que le
    commandant n'a pas signé, le document porte la mention `DRAFT` et ne revendique
    **aucun** original signé. La version booking écrivait « Trois originaux signés
    (3 OBL) » **inconditionnellement**, y compris sur un document que personne n'avait
    signé — et c'est le format **éditable**, donc celui qui circule le plus.
    """
    from app.services.bl_workflow import FROZEN_STATES
    from app.templating import brand_for_lang

    brand: dict[str, Any] = brand_for_lang("fr")
    issued = issued_at or datetime.now(UTC)
    is_signed = getattr(batch, "bl_state", None) in FROZEN_STATES
    doc = _new_document()

    _title(doc, "BILL OF LADING · CONNAISSEMENT")
    _ref_centered(doc, bl_number)
    if not is_signed:
        warn = doc.add_paragraph()
        run = warn.add_run(
            "PROJET — SANS VALEUR DE TITRE. Ce document n'a pas été signé par le "
            "commandant : il ne constitue ni un original négociable, ni une preuve de "
            "mise à bord, et ne vaut pas remise de la marchandise."
        )
        run.bold = True
        run.italic = True
    doc.add_paragraph()

    parties = doc.add_table(rows=2, cols=2)
    parties.style = "Table Grid"
    shipper = "\n".join(
        line
        for line in [
            batch.shipper_name,
            batch.shipper_address,
            f"{batch.shipper_postal or ''} {batch.shipper_city or ''}".strip(),
            batch.shipper_country,
        ]
        if line
    )
    carrier = "\n".join(
        [brand["raison_sociale"], brand["adresse"], brand["telephone"], brand["email"]]
    )
    consignee = "\n".join(
        line
        for line in [
            batch.consignee_name,
            batch.consignee_address,
            f"{batch.consignee_postal or ''} {batch.consignee_city or ''}".strip(),
            batch.consignee_country,
        ]
        if line
    )
    notify = "\n".join(
        line
        for line in [
            batch.notify_name,
            batch.notify_address,
            f"{batch.notify_postal or ''} {batch.notify_city or ''}".strip(),
            batch.notify_country,
        ]
        if line
    )
    parties.rows[0].cells[0].text = "Shipper · Expéditeur\n" + (shipper or "—")
    parties.rows[0].cells[1].text = "Carrier · Transporteur\n" + carrier
    parties.rows[1].cells[0].text = "Consignee · Destinataire\n" + (consignee or "—")
    parties.rows[1].cells[1].text = "Notify party\n" + (notify or "—")
    doc.add_paragraph()

    _section(doc, "Voyage")
    rows: list[tuple[str, str]] = []
    if vessel is not None:
        desc = vessel.name
        extra = [
            x
            for x in (
                f"IMO {vessel.imo_number}" if getattr(vessel, "imo_number", None) else None,
                getattr(vessel, "flag", None),
            )
            if x
        ]
        desc += f" ({vessel.code}" + ((" · " + " · ".join(extra)) if extra else "") + ")"
        rows.append(("Navire / Vessel", desc))
    if leg is not None:
        rows.append(("Voyage · Leg code", leg.leg_code))
        rows.append(("ETD", _fmt_date(leg.etd, "%d/%m/%Y %H:%M UTC")))
        rows.append(("ETA", _fmt_date(leg.eta, "%d/%m/%Y %H:%M UTC")))
    if pol is not None:
        rows.append(("Port of Loading (POL)", f"{pol.name} ({pol.locode} · {pol.country})"))
    if pod is not None:
        rows.append(("Port of Discharge (POD)", f"{pod.name} ({pod.locode} · {pod.country})"))
    # §5.0 — omise plutôt qu'inventée : une date de mise à bord fausse (a fortiori
    # future) est une fraude documentaire.
    sob_value = getattr(shipped_on_board, "value", None) if shipped_on_board else None
    rows.append(
        ("Shipped on board", _fmt_date(sob_value, "%d/%m/%Y") if sob_value else "— non constatée")
    )
    _kv_table(doc, rows)
    doc.add_paragraph()

    _section(doc, "Goods · Marchandises")
    goods = doc.add_table(rows=1, cols=6)
    goods.style = "Table Grid"
    for idx, label in enumerate(
        ["Format", "Qté palettes", "Description", "HS", "Poids (kg)", "IMDG"]
    ):
        goods.rows[0].cells[idx].text = label
        if goods.rows[0].cells[idx].paragraphs[0].runs:
            goods.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
    cells = goods.add_row().cells
    cells[0].text = batch.pallet_format or "—"
    cells[1].text = str(batch.pallet_count or 0)
    cells[2].text = batch.description_of_goods or batch.type_of_goods or "—"
    cells[3].text = batch.hs_code or "—"
    cells[4].text = str(batch.weight_kg) if batch.weight_kg is not None else "—"
    if getattr(batch, "hazardous", False):
        imdg = batch.imdg_class or "IMDG"
        if getattr(batch, "un_number", None):
            imdg += f" · UN {batch.un_number}"
        cells[5].text = imdg
    else:
        cells[5].text = "—"
    doc.add_paragraph()
    if batch.marks_and_numbers:
        doc.add_paragraph(f"Marks & numbers : {batch.marks_and_numbers}")

    _section(doc, "Conditions")
    doc.add_paragraph(
        "Transport assuré conformément aux Règles de La Haye-Visby. La responsabilité "
        "du transporteur est plafonnée selon les conventions internationales en vigueur."
    )

    doc.add_paragraph()
    stamp = doc.add_paragraph()
    stamp.add_run(f"Émis{f' à {pol.name}' if pol is not None else ''} le {_fmt_date(issued)}\n")
    if is_signed:
        stamp.add_run("Trois originaux signés (3 OBL)\n").italic = True
        stamp.add_run(f"Signé par : {batch.bl_signed_by_name or '—'}\n")
        if batch.bl_signed_at:
            stamp.add_run(f"Le : {_fmt_date(batch.bl_signed_at, '%d/%m/%Y %H:%M UTC')}\n")
        if batch.bl_signature_hash:
            stamp.add_run(f"Empreinte SHA-256 : {batch.bl_signature_hash}\n")
        stamp.add_run("\nCachet et signature du transporteur")
    else:
        stamp.add_run("Aucun original signé à ce stade. À la signature : 3 originaux.\n").italic = (
            True
        )

    _footer(doc)
    suffix = "" if is_signed else "-DRAFT"
    return _serialize(doc, f"{bl_number}{suffix}.docx")


# ---------------------------------------------------------------------------
# Booking note (contrat de réservation d'espace en cale)
# ---------------------------------------------------------------------------


def _bn_cell(cell, label: str, value: str | None) -> None:
    """Case du formulaire : intitulé en petit, valeur en dessous.

    Reproduit la présentation du gabarit BIMCO, où chaque case porte son
    intitulé imprimé et la valeur saisie en dessous. Une valeur absente laisse
    la case **visiblement vide** (un tiret) : elle est à compléter, pas à
    deviner.
    """
    from docx.shared import Pt

    cell.text = ""
    head = cell.paragraphs[0]
    head_run = head.add_run(label)
    head_run.font.size = Pt(7)
    head_run.bold = True
    head_run.font.color.rgb = _teal_color()

    body = cell.add_paragraph()
    body_run = body.add_run(value if value else "—")
    body_run.font.size = Pt(9)
    if value:
        body_run.bold = True


def build_booking_note_docx(*, note, offer, leg=None) -> DocxBytes:
    """Booking note Word depuis un ``BookingNote`` (trame CONLINEBOOKING).

    Le recto reprend les cases du gabarit fourni par la direction ; le verso
    imprime les conditions générales verbatim (``booking_note_terms``). Le
    document est construit programmatiquement avec ``python-docx`` : aucun
    moteur de gabarit n'intervient, donc aucun texte saisi ne peut être
    interprété comme de la mise en forme.

    Un brouillon porte un **filigrane textuel** en tête : il circule parfois en
    interne avant diffusion, et rien ne doit laisser croire qu'il engage déjà.
    """
    from docx.shared import Pt

    from app.services.booking_note_terms import (
        BOOKING_NOTE_TERMS,
        BOOKING_NOTE_TERMS_TITLE,
    )

    doc = _new_document()

    _title(doc, f"BOOKING NOTE NUMBER – {note.reference}")
    if note.status != "diffusee":
        draft = doc.add_paragraph()
        draft_run = draft.add_run(
            "BROUILLON — document de travail, non diffusé au client"
        )
        draft_run.bold = True
        draft_run.font.size = Pt(10)
        from docx.shared import RGBColor

        draft_run.font.color.rgb = RGBColor(0xB4, 0x71, 0x48)  # cuivre NEWTOWT

    place_date = ", ".join(
        bit for bit in (note.issue_place, _fmt_date(note.issued_on)) if bit and bit != "—"
    )

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    def _row(left: tuple[str, str | None], right: tuple[str, str | None]) -> None:
        cells = table.add_row().cells
        _bn_cell(cells[0], left[0], left[1])
        _bn_cell(cells[1], right[0], right[1])

    _row(("Agents", note.agents_pod), ("Place and date", place_date or None))
    _row(("Carrier", "NEWTOWT\n52 Quai Frissard\n76600 LE HAVRE\nFRANCE"), ("Vessel", note.vessel_name))
    _row(
        (
            "Merchant*",
            "\n".join(
                bit
                for bit in (
                    note.merchant_name,
                    note.merchant_contact,
                    note.merchant_address,
                    note.merchant_email,
                )
                if bit
            )
            or None,
        ),
        ("Time for shipment (about)", note.time_for_shipment),
    )
    _row(("Port of loading**", note.pol_text), ("Port of discharge**", note.pod_text))
    _row(
        (
            "Merchant's representatives at loading port",
            "TO BE CONFIRMED at time of shipment of the Bill of lading",
        ),
        (
            "Container No./Seal No./Marks and Nos. (if available)",
            "TO BE CONFIRMED at time of shipment of the Bill of lading",
        ),
    )
    _row(
        ("Number and kind of packages, description of cargo", note.cargo_description),
        ("Gross weight, kg (if available)", "1000 kg maximum per pallet"),
    )
    _row(
        ("Freight details and charges", note.freight_terms),
        ("Special terms, if agreed", note.special_terms),
    )
    _row(
        (
            "Freight (state pre-payable or payable at destination)\n"
            "Subject to deadfreight (100%) if shipment canceled 3 months before ETD",
            note.payment_terms,
        ),
        (
            "Measurement, m3 (if available)",
            "TO BE CONFIRMED at time of shipment of the Bill of lading",
        ),
    )

    doc.add_paragraph()
    agreement = doc.add_paragraph(
        "It is hereby agreed that this Contract shall be performed subject to the terms "
        "contained on Page 1 and 2 hereof which shall prevail over any previous arrangements, "
        "and which shall in turn be superseded (except as to dead freight) by the terms of the "
        "Bill of Lading."
    )
    for run in agreement.runs:
        run.font.size = Pt(8)

    sign = doc.add_table(rows=2, cols=2)
    sign.style = "Table Grid"
    sign.rows[0].cells[0].text = "Signature (Merchant)"
    sign.rows[0].cells[1].text = "Signature (Carrier)"
    sign.rows[1].cells[0].text = "\n\n"
    sign.rows[1].cells[1].text = "\n\n"

    for footnote in (
        "* As defined hereinafter (Cl. 1)",
        "** (or so near thereunto as the Vessel may safely get and lie always afloat)",
    ):
        para = doc.add_paragraph(footnote)
        for run in para.runs:
            run.font.size = Pt(7)

    # ── Verso : conditions générales, verbatim ───────────────────────────
    from docx.enum.text import WD_BREAK

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    _section(doc, BOOKING_NOTE_TERMS_TITLE)
    for clause in BOOKING_NOTE_TERMS:
        para = doc.add_paragraph(clause)
        for run in para.runs:
            run.font.size = Pt(7)

    _footer(doc)
    return _serialize(doc, f"Booking_Note_{note.reference}.docx")
