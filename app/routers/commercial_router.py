"""Commercial — clients (FF/Shipper), grilles tarifaires, offres, commandes.

Reprises de la V3.0.0 :
- Brackets dégressifs par volume (lt50→full ship).
- Génération de référence : ORD-YYYY-NNNN, RG-YYYY-NNNN, RO-YYYY-NNNN.
- Conversion offre → commande.
- Audit trail activity_logs sur toutes les actions.
"""

from __future__ import annotations

import io
import json
from datetime import UTC, datetime
from datetime import date as _date
from decimal import Decimal, InvalidOperation

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, Depends, Form, HTTPException, Request
from fastapi.responses import HTMLResponse, RedirectResponse, Response
from sqlalchemy import or_, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.database import get_db
from app.models.client_account import ClientAccount
from app.models.commercial import (
    CLIENT_TYPES,
    RATE_OPTION_UNIT_LABELS,
    RATE_OPTION_UNITS,
    Client,
    Order,
    OrderAssignment,
    RateGrid,
    RateGridLine,
    RateGridOption,
    RateOffer,
)
from app.models.leg import Leg
from app.models.port import Port
from app.models.quote import Quote
from app.models.vessel import Vessel
from app.permissions import require_permission
from app.services import capacity as capacity_svc
from app.services.activity import record as activity_record
from app.services.commercial import (
    bracket_rate,
    compatible_legs_for_order,
    default_brackets_for,
    leg_is_late_for_order,
    next_grid_reference,
    next_offer_reference,
    next_order_reference,
    pick_bracket,
    suggest_leg_for_order,
)
from app.services.quoting import (
    _match_route,
    backfill_default_grids,
    compute_route_economics,
)
from app.templating import templates

router = APIRouter(prefix="/commercial", tags=["commercial"])


def _hx_or_redirect(request: Request, target: str):
    """303 classique, ou header HX-Redirect si la requête vient d'HTMX."""
    if request.headers.get("hx-request"):
        return Response(status_code=200, headers={"HX-Redirect": target})
    return RedirectResponse(url=target, status_code=303)


# ────────────────────────────────────────────── Landing
@router.get("", response_class=HTMLResponse)
@router.get("/", response_class=HTMLResponse)
async def commercial_index(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    clients_count = (await db.scalar(select(__import__("sqlalchemy").func.count(Client.id)))) or 0
    grids_active = (
        await db.scalar(
            select(__import__("sqlalchemy").func.count(RateGrid.id)).where(
                RateGrid.status == "active"
            )
        )
    ) or 0
    offers_open = (
        await db.scalar(
            select(__import__("sqlalchemy").func.count(RateOffer.id)).where(
                RateOffer.status.in_(("draft", "sent"))
            )
        )
    ) or 0
    orders_open = (
        await db.scalar(
            select(__import__("sqlalchemy").func.count(Order.id)).where(
                Order.status.in_(("draft", "confirmed", "loaded"))
            )
        )
    ) or 0
    last_orders = list(
        (
            await db.execute(
                select(Order)
                .options(selectinload(Order.client))
                .order_by(Order.created_at.desc())
                .limit(10)
            )
        )
        .scalars()
        .all()
    )

    # Remplissage des legs en cours de commercialisation : legs réservables, non
    # encore appareillés, dont la fenêtre de réservation est ouverte. Le service
    # capacity lève NotBookable/BookingClosed pour les legs hors commercialisation
    # → on s'en sert comme filtre. Occupation = (bookings + commandes) / capacité.
    fill_legs, fill_summary = await _commercialization_fill(db)

    return templates.TemplateResponse(
        "staff/commercial/index.html",
        {
            "request": request,
            "user": user,
            "clients_count": clients_count,
            "grids_active": grids_active,
            "offers_open": offers_open,
            "orders_open": orders_open,
            "last_orders": last_orders,
            "fill_legs": fill_legs,
            "fill_summary": fill_summary,
        },
    )


async def _commercialization_fill(db: AsyncSession) -> tuple[list[dict], dict]:
    """Taux de remplissage des legs en cours de commercialisation.

    Renvoie ``(lignes, synthèse)`` où chaque ligne décrit un leg réservable
    (code, navire, ports, ETD, capacité, réservé, disponible, % occupation) et
    la synthèse agrège capacité/réservé sur l'ensemble.
    """
    now = datetime.now(UTC)
    candidate_legs = list(
        (
            await db.execute(
                select(Leg)
                .where(Leg.is_bookable.is_(True))
                .where(Leg.atd.is_(None))
                .where(Leg.etd >= now)
                .order_by(Leg.etd.asc())
            )
        )
        .scalars()
        .all()
    )
    if not candidate_legs:
        return [], {"capacity": 0, "reserved": 0, "occupancy_pct": 0.0, "count": 0}

    vessels = {
        v.id: v for v in (await db.execute(select(Vessel))).scalars().all()
    }
    port_ids = {leg.departure_port_id for leg in candidate_legs} | {
        leg.arrival_port_id for leg in candidate_legs
    }
    ports = {
        p.id: p
        for p in (await db.execute(select(Port).where(Port.id.in_(port_ids)))).scalars().all()
    }

    rows: list[dict] = []
    total_cap = 0
    total_res = 0
    for leg in candidate_legs:
        try:
            info = await capacity_svc.get_available_capacity(db, leg.id)
        except (capacity_svc.NotBookable, capacity_svc.BookingClosed):
            continue
        except ValueError:
            continue
        vessel = vessels.get(leg.vessel_id)
        pol = ports.get(leg.departure_port_id)
        pod = ports.get(leg.arrival_port_id)
        total_cap += info.capacity_palettes
        total_res += info.reserved_palettes
        rows.append(
            {
                "leg_id": leg.id,
                "leg_code": leg.leg_code,
                "vessel_name": vessel.name if vessel else "—",
                "pol_locode": pol.locode if pol else "?",
                "pod_locode": pod.locode if pod else "?",
                "etd": leg.etd,
                "booking_close_at": leg.booking_close_at,
                "capacity": info.capacity_palettes,
                "reserved": info.reserved_palettes,
                "available": info.available_palettes,
                "occupancy_pct": info.occupancy_pct,
            }
        )
    summary = {
        "capacity": total_cap,
        "reserved": total_res,
        "occupancy_pct": round(100 * total_res / total_cap, 1) if total_cap else 0.0,
        "count": len(rows),
    }
    return rows, summary


# ────────────────────────────────────────────── Clients (FF / Shipper)
@router.get("/clients", response_class=HTMLResponse)
async def clients_list(
    request: Request,
    q: str | None = None,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    from app.services.pipedrive_sync import is_configured

    # Recherche base client : nom, contact, e-mail, n° TVA, ville/pays.
    query = select(Client)
    term = (q or "").strip()
    if term:
        like = f"%{term}%"
        query = query.where(
            or_(
                Client.name.ilike(like),
                Client.contact_name.ilike(like),
                Client.contact_email.ilike(like),
                Client.vat_number.ilike(like),
                Client.country.ilike(like),
            )
        )
    clients = list((await db.execute(query.order_by(Client.name.asc()))).scalars().all())
    return templates.TemplateResponse(
        "staff/commercial/clients.html",
        {
            "request": request,
            "user": user,
            "clients": clients,
            "types": CLIENT_TYPES,
            "pipedrive_configured": is_configured(),
            "search_q": term,
        },
    )


@router.post("/clients/sync")
async def clients_sync_pipedrive(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
) -> RedirectResponse:
    """Remonte les organisations Pipedrive dans la liste des clients."""
    from app.services.pipedrive_sync import sync_clients

    result = await sync_clients(db)
    if not result["configured"]:
        return RedirectResponse(url="/commercial/clients?pd=disabled", status_code=303)
    await activity_record(
        db,
        action="sync",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="client",
        entity_label=f"pipedrive: +{result['created']} / ~{result['updated']}",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(
        url=(
            f"/commercial/clients?pd=ok&created={result['created']}"
            f"&updated={result['updated']}&skipped={result.get('skipped', 0)}"
            f"&linked={result.get('linked', 0)}"
        ),
        status_code=303,
    )


@router.post("/clients")
async def client_create(
    request: Request,
    name: str = Form(...),
    client_type: str = Form(...),
    contact_name: str | None = Form(None),
    contact_email: str | None = Form(None),
    contact_phone: str | None = Form(None),
    phone_dial_code: str | None = Form(None),
    address: str | None = Form(None),
    country: str | None = Form(None),
    vat_number: str | None = Form(None),
    notes: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    if client_type not in CLIENT_TYPES:
        raise HTTPException(status_code=400, detail="invalid client_type")
    c = Client(
        name=name.strip(),
        client_type=client_type,
        contact_name=(contact_name or "").strip() or None,
        contact_email=(contact_email or "").strip() or None,
        contact_phone=_compose_phone(phone_dial_code, contact_phone),
        address=(address or "").strip() or None,
        country=(country or "").strip().upper()[:2] or None,
        vat_number=(vat_number or "").strip() or None,
        notes=notes,
    )
    db.add(c)
    await db.flush()
    await activity_record(
        db,
        action="create",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="client",
        entity_id=c.id,
        entity_label=c.name,
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url="/commercial/clients", status_code=303)


@router.get("/clients/{client_id}", response_class=HTMLResponse)
async def client_detail(
    client_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    client = await db.get(Client, client_id)
    if client is None:
        raise HTTPException(status_code=404)
    grids = list(
        (
            await db.execute(
                select(RateGrid)
                .options(selectinload(RateGrid.lines))
                .where(RateGrid.client_id == client.id)
                .order_by(RateGrid.created_at.desc())
            )
        )
        .scalars()
        .all()
    )
    linked_accounts = list(
        (
            await db.execute(
                select(ClientAccount)
                .where(ClientAccount.commercial_client_id == client.id)
                .order_by(ClientAccount.email.asc())
            )
        )
        .scalars()
        .all()
    )
    unlinked_accounts = list(
        (
            await db.execute(
                select(ClientAccount)
                .where(ClientAccount.commercial_client_id.is_(None))
                .order_by(ClientAccount.email.asc())
                .limit(200)
            )
        )
        .scalars()
        .all()
    )
    return templates.TemplateResponse(
        "staff/commercial/client_detail.html",
        {
            "request": request,
            "user": user,
            "client": client,
            "grids": grids,
            "linked_accounts": linked_accounts,
            "unlinked_accounts": unlinked_accounts,
        },
    )


@router.get("/clients/{client_id}/edit", response_class=HTMLResponse)
async def client_edit_form(
    client_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
) -> HTMLResponse:
    client = await db.get(Client, client_id)
    if client is None:
        raise HTTPException(status_code=404)
    return templates.TemplateResponse(
        "staff/commercial/client_form.html",
        {"request": request, "user": user, "client": client, "types": CLIENT_TYPES},
    )


@router.post("/clients/{client_id}/edit")
async def client_edit(
    client_id: int,
    request: Request,
    name: str = Form(...),
    client_type: str = Form(...),
    contact_name: str | None = Form(None),
    contact_email: str | None = Form(None),
    contact_phone: str | None = Form(None),
    phone_dial_code: str | None = Form(None),
    address: str | None = Form(None),
    country: str | None = Form(None),
    vat_number: str | None = Form(None),
    notes: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    client = await db.get(Client, client_id)
    if client is None:
        raise HTTPException(status_code=404)
    if client_type not in CLIENT_TYPES:
        raise HTTPException(status_code=400, detail="invalid client_type")
    client.name = name.strip()
    client.client_type = client_type
    client.contact_name = (contact_name or "").strip() or None
    client.contact_email = (contact_email or "").strip() or None
    client.contact_phone = _compose_phone(phone_dial_code, contact_phone)
    client.address = (address or "").strip() or None
    client.country = (country or "").strip().upper()[:2] or None
    client.vat_number = (vat_number or "").strip() or None
    client.notes = (notes or "").strip() or None
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="client",
        entity_id=client.id,
        entity_label=client.name,
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/clients/{client.id}")


@router.post("/clients/{client_id}/toggle-active")
async def client_toggle_active(
    client_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    client = await db.get(Client, client_id)
    if client is None:
        raise HTTPException(status_code=404)
    client.is_active = not client.is_active
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="client",
        entity_id=client.id,
        entity_label=client.name,
        detail="réactivé" if client.is_active else "désactivé",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/clients/{client.id}")


@router.post("/clients/{client_id}/accounts/link")
async def client_account_link(
    client_id: int,
    request: Request,
    account_id: int = Form(...),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    client = await db.get(Client, client_id)
    if client is None:
        raise HTTPException(status_code=404, detail="client introuvable")
    account = await db.get(ClientAccount, account_id)
    if account is None:
        raise HTTPException(status_code=404, detail="compte plateforme introuvable")
    if account.commercial_client_id is not None and account.commercial_client_id != client.id:
        raise HTTPException(status_code=400, detail="compte déjà relié à un autre client")
    account.commercial_client_id = client.id
    await db.flush()
    await activity_record(
        db,
        action="account_link",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="client",
        entity_id=client.id,
        entity_label=client.name,
        detail=f"compte plateforme #{account.id} relié",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/clients/{client.id}")


@router.post("/clients/{client_id}/accounts/{account_id}/unlink")
async def client_account_unlink(
    client_id: int,
    account_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    client = await db.get(Client, client_id)
    if client is None:
        raise HTTPException(status_code=404, detail="client introuvable")
    account = await db.get(ClientAccount, account_id)
    if account is None or account.commercial_client_id != client.id:
        raise HTTPException(status_code=404, detail="compte plateforme non relié à ce client")
    account.commercial_client_id = None
    await db.flush()
    await activity_record(
        db,
        action="account_unlink",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="client",
        entity_id=client.id,
        entity_label=client.name,
        detail=f"compte plateforme #{account.id} délié",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/clients/{client.id}")


# ────────────────────────────────────────────── Rate grids
def _compose_phone(dial_code: str | None, number: str | None) -> str | None:
    """Combine indicatif international + numéro local en un téléphone unique.

    - numéro vide ⇒ None ;
    - numéro déjà en format international (commence par ``+`` ou ``00``) ⇒
      conservé tel quel (l'indicatif est ignoré pour éviter un doublon) ;
    - sinon ⇒ ``{indicatif} {numéro}`` (ex. ``+33 6 12 34 56 78``).
    """
    num = (number or "").strip()
    if not num:
        return None
    if num.startswith("+") or num.startswith("00"):
        return num[:50]
    code = (dial_code or "").strip()
    composed = f"{code} {num}".strip() if code else num
    return composed[:50]


def _opt_decimal(value: str | None) -> Decimal | None:
    """Parse un décimal optionnel (champ vide → None ; invalide → None)."""
    raw = (value or "").strip().replace(",", ".")
    if not raw:
        return None
    try:
        return Decimal(raw)
    except (InvalidOperation, ValueError):
        return None


def _opt_int(value: str | None) -> int | None:
    """Parse un entier optionnel (champ vide / invalide → None)."""
    raw = (value or "").strip()
    if not raw:
        return None
    try:
        return int(raw)
    except (TypeError, ValueError):
        return None


def _opt_date(value: str | None, field: str = "date") -> _date | None:
    """Parse une date ISO optionnelle (``YYYY-MM-DD``) ; vide → None.

    Une saisie non vide mais invalide lève un 400 (cohérent avec
    ``_clean_locode``) plutôt que d'être silencieusement ignorée — sinon une
    fenêtre de livraison mal saisie désactiverait sans bruit l'alerte « hors
    délai » qui est le cœur de COM-01.
    """
    raw = (value or "").strip()
    if not raw:
        return None
    try:
        return _date.fromisoformat(raw[:10])
    except ValueError as exc:
        raise HTTPException(
            status_code=400, detail=f"{field} invalide : date au format AAAA-MM-JJ"
        ) from exc


async def _route_ports(db: AsyncSession) -> list[Port]:
    """Ports desservis par au moins un leg (pour pré-remplir POL/POD des grilles)."""
    ids = select(Leg.departure_port_id).union(select(Leg.arrival_port_id))
    return list(
        (await db.execute(select(Port).where(Port.id.in_(ids)).order_by(Port.name))).scalars().all()
    )


def _clean_locode(value: str | None, field: str) -> str | None:
    """LOCODE UN normalisé (majuscules) ou None ; 400 si format invalide."""
    code = (value or "").strip().upper()
    if not code:
        return None
    if len(code) != 5 or not code.isalnum():
        raise HTTPException(
            status_code=400,
            detail=f"{field} invalide : LOCODE UN sur 5 caractères (ex. FRLEH)",
        )
    return code


def _validate_grid_header(*, client_id: int | None, is_default: bool) -> None:
    """Règles d'en-tête d'une grille (les routes sont des lignes séparées).

    - grille par défaut : aucun client ;
    - grille non défaut : client obligatoire.
    """
    if is_default and client_id is not None:
        raise HTTPException(
            status_code=400,
            detail="une grille par défaut ne peut pas être rattachée à un client",
        )
    if not is_default and client_id is None:
        raise HTTPException(
            status_code=400,
            detail="une grille doit être rattachée à un client, ou marquée « grille par défaut »",
        )


async def _vessels(db: AsyncSession) -> list[Vessel]:
    """Navires actifs (pour le select navire de référence d'une grille)."""
    return list(
        (await db.execute(select(Vessel).where(Vessel.is_active.is_(True)).order_by(Vessel.name)))
        .scalars()
        .all()
    )


async def _grid_editable(db: AsyncSession, grid_id: int) -> RateGrid:
    """Grille (avec routes) éditable : 404 si absente, 400 si active (verrouillée)."""
    grid = (
        await db.execute(
            select(RateGrid).options(selectinload(RateGrid.lines)).where(RateGrid.id == grid_id)
        )
    ).scalar_one_or_none()
    if grid is None:
        raise HTTPException(status_code=404)
    if grid.status == "active":
        raise HTTPException(
            status_code=400,
            detail="Grille active verrouillée — repassez-la en brouillon pour la modifier.",
        )
    return grid


@router.get("/grids", response_class=HTMLResponse)
async def grids_list(
    request: Request,
    created: int | None = None,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    grids = list(
        (
            await db.execute(
                select(RateGrid)
                .options(selectinload(RateGrid.client), selectinload(RateGrid.lines))
                .order_by(RateGrid.created_at.desc())
            )
        )
        .scalars()
        .all()
    )
    return templates.TemplateResponse(
        "staff/commercial/grids.html",
        {"request": request, "user": user, "grids": grids, "created": created},
    )


@router.post("/grids/backfill-defaults")
async def grids_backfill_defaults(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    """Crée la grille par défaut de chaque route POL/POD du planning."""
    count = await backfill_default_grids(db)
    await db.flush()
    await activity_record(
        db,
        action="grids_backfill_defaults",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid",
        detail=f"{count} grilles créées",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/grids?created={count}")


@router.get("/grids/new", response_class=HTMLResponse)
async def grid_new_form(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
) -> HTMLResponse:
    clients = list(
        (await db.execute(select(Client).where(Client.is_active.is_(True)).order_by(Client.name)))
        .scalars()
        .all()
    )
    return templates.TemplateResponse(
        "staff/commercial/grid_form.html",
        {
            "request": request,
            "user": user,
            "clients": clients,
            "vessels": await _vessels(db),
            "grid": None,
        },
    )


@router.post("/grids")
async def grid_create(
    request: Request,
    client_id: int | None = Form(None),
    vessel_id: int | None = Form(None),
    is_default: bool = Form(False),
    valid_from: str = Form(...),
    valid_to: str | None = Form(None),
    adjustment_index: float = Form(1.0),
    bl_fee: str | None = Form(None),
    booking_fee: str | None = Form(None),
    hazardous_surcharge_pct: str | None = Form(None),
    min_charge_eur: str | None = Form(None),
    volume_commitment: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    _validate_grid_header(client_id=client_id, is_default=is_default)
    client = None
    if client_id is not None:
        client = await db.get(Client, client_id)
        if client is None:
            raise HTTPException(status_code=404, detail="client introuvable")
    if vessel_id is not None and await db.get(Vessel, vessel_id) is None:
        raise HTTPException(status_code=404, detail="navire introuvable")
    ref = await next_grid_reference(db)
    # Brackets de volume au niveau grille selon le type de client (shipper :
    # dégressif complet ; FF : tarif flat). Stockés en JSON sur l'en-tête.
    brackets = default_brackets_for(client.client_type if client else "shipper")
    grid = RateGrid(
        reference=ref,
        client_id=client.id if client else None,
        vessel_id=vessel_id,
        is_default=is_default,
        status="draft",
        valid_from=_date.fromisoformat(valid_from),
        valid_to=_date.fromisoformat(valid_to) if valid_to else None,
        currency="EUR",
        adjustment_index=Decimal(str(adjustment_index)),
        bl_fee=_opt_decimal(bl_fee),
        booking_fee=_opt_decimal(booking_fee),
        hazardous_surcharge_pct=_opt_decimal(hazardous_surcharge_pct),
        min_charge_eur=_opt_decimal(min_charge_eur),
        volume_commitment=_opt_int(volume_commitment),
        brackets_json=json.dumps(brackets),
    )
    db.add(grid)
    await db.flush()
    await activity_record(
        db,
        action="create",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid",
        entity_id=grid.id,
        entity_label=ref,
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/grids/{grid.id}", status_code=303)


@router.get("/grids/{grid_id}/edit", response_class=HTMLResponse)
async def grid_edit_form(
    grid_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
) -> HTMLResponse:
    grid = await db.get(RateGrid, grid_id)
    if grid is None:
        raise HTTPException(status_code=404)
    clients = list(
        (await db.execute(select(Client).where(Client.is_active.is_(True)).order_by(Client.name)))
        .scalars()
        .all()
    )
    return templates.TemplateResponse(
        "staff/commercial/grid_form.html",
        {
            "request": request,
            "user": user,
            "clients": clients,
            "vessels": await _vessels(db),
            "grid": grid,
        },
    )


@router.post("/grids/{grid_id}/edit")
async def grid_edit(
    grid_id: int,
    request: Request,
    client_id: int | None = Form(None),
    vessel_id: int | None = Form(None),
    is_default: bool = Form(False),
    valid_from: str = Form(...),
    valid_to: str | None = Form(None),
    adjustment_index: float = Form(1.0),
    bl_fee: str | None = Form(None),
    booking_fee: str | None = Form(None),
    hazardous_surcharge_pct: str | None = Form(None),
    min_charge_eur: str | None = Form(None),
    volume_commitment: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    grid = await db.get(RateGrid, grid_id)
    if grid is None:
        raise HTTPException(status_code=404)
    # Cycle de vie : une grille active est verrouillée (cf. Module 6) — il faut
    # la repasser en brouillon (ou en créer une nouvelle) avant de l'éditer.
    if grid.status == "active":
        raise HTTPException(
            status_code=400,
            detail="Grille active verrouillée — repassez-la en brouillon pour la modifier.",
        )
    _validate_grid_header(client_id=client_id, is_default=is_default)
    if client_id is not None and await db.get(Client, client_id) is None:
        raise HTTPException(status_code=404, detail="client introuvable")
    if vessel_id is not None and await db.get(Vessel, vessel_id) is None:
        raise HTTPException(status_code=404, detail="navire introuvable")
    grid.client_id = client_id
    grid.vessel_id = vessel_id
    grid.is_default = is_default
    grid.valid_from = _date.fromisoformat(valid_from)
    grid.valid_to = _date.fromisoformat(valid_to) if valid_to else None
    grid.adjustment_index = Decimal(str(adjustment_index))
    grid.bl_fee = _opt_decimal(bl_fee)
    grid.booking_fee = _opt_decimal(booking_fee)
    grid.hazardous_surcharge_pct = _opt_decimal(hazardous_surcharge_pct)
    grid.min_charge_eur = _opt_decimal(min_charge_eur)
    grid.volume_commitment = _opt_int(volume_commitment)
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid",
        entity_id=grid.id,
        entity_label=grid.reference,
        detail="edited",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/grids/{grid.id}")


@router.get("/grids/{grid_id}", response_class=HTMLResponse)
async def grid_detail(
    grid_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    grid = (
        await db.execute(
            select(RateGrid)
            .options(
                selectinload(RateGrid.client),
                selectinload(RateGrid.lines),
                selectinload(RateGrid.options),
            )
            .where(RateGrid.id == grid_id)
        )
    ).scalar_one_or_none()
    if grid is None:
        raise HTTPException(status_code=404)
    vessel = await db.get(Vessel, grid.vessel_id) if grid.vessel_id else None
    return templates.TemplateResponse(
        "staff/commercial/grid_detail.html",
        {
            "request": request,
            "user": user,
            "grid": grid,
            "vessel": vessel,
            "route_ports": await _route_ports(db),
            "option_units": RATE_OPTION_UNITS,
            "option_unit_labels": RATE_OPTION_UNIT_LABELS,
        },
    )


@router.post("/grids/{grid_id}/activate")
async def grid_activate(
    grid_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    grid = await db.get(RateGrid, grid_id)
    if grid is None:
        raise HTTPException(status_code=404)
    # Une seule grille active par périmètre : par client pour une grille client,
    # une seule grille par défaut active globalement. Les autres actives du même
    # périmètre sont marquées « superseded ».
    others = select(RateGrid).where(RateGrid.id != grid.id, RateGrid.status == "active")
    if grid.client_id is not None:
        others = others.where(RateGrid.client_id == grid.client_id)
    else:
        others = others.where(RateGrid.client_id.is_(None), RateGrid.is_default.is_(True))
    superseded = 0
    for other in (await db.execute(others)).scalars().all():
        other.status = "superseded"
        superseded += 1
    grid.status = "active"
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid",
        entity_id=grid.id,
        entity_label=grid.reference,
        detail=f"activated (superseded={superseded})",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/grids/{grid_id}", status_code=303)


@router.post("/grids/{grid_id}/draft")
async def grid_set_draft(
    grid_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    """Repasse une grille en brouillon (déverrouille l'édition)."""
    grid = await db.get(RateGrid, grid_id)
    if grid is None:
        raise HTTPException(status_code=404)
    grid.status = "draft"
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid",
        entity_id=grid.id,
        entity_label=grid.reference,
        detail="set_draft",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/grids/{grid_id}", status_code=303)


async def _recalc_route(db: AsyncSession, grid: RateGrid, route: RateGridLine) -> None:
    """Recalcule distance/nav_days/opex/base_rate d'une route (OPEX du navire)."""
    leg = await db.get(Leg, route.leg_id) if route.leg_id else None
    distance, nav_days, opex_daily, base = await compute_route_economics(
        db,
        pol_locode=route.pol_locode,
        pod_locode=route.pod_locode,
        vessel_id=grid.vessel_id,
        leg=leg,
    )
    route.distance_nm = distance
    route.nav_days = nav_days
    route.opex_daily = opex_daily
    route.base_rate = base
    route.is_manual = False


@router.post("/grids/{grid_id}/recalculate")
async def grid_recalculate(
    grid_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    """Recalcule le base_rate de toutes les routes non-manuelles (OPEX × jours / 850)."""
    grid = (
        await db.execute(
            select(RateGrid).options(selectinload(RateGrid.lines)).where(RateGrid.id == grid_id)
        )
    ).scalar_one_or_none()
    if grid is None:
        raise HTTPException(status_code=404)
    if grid.status == "active":
        raise HTTPException(
            status_code=400,
            detail="Grille active verrouillée — repassez-la en brouillon pour la recalculer.",
        )
    recalculated = 0
    for route in grid.lines:
        if route.is_manual:
            continue
        await _recalc_route(db, grid, route)
        recalculated += 1
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid",
        entity_id=grid.id,
        entity_label=grid.reference,
        detail=f"recalculated routes={recalculated}",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/grids/{grid_id}", status_code=303)


# ────────────────────────────────────────────── Grid routes (lignes-routes)
@router.post("/grids/{grid_id}/routes")
async def grid_route_create(
    grid_id: int,
    request: Request,
    pol_locode: str = Form(...),
    pod_locode: str = Form(...),
    distance_nm: str | None = Form(None),
    base_rate: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    """Ajoute une route POL→POD à la grille (distance saisie ou résolue)."""
    grid = await _grid_editable(db, grid_id)
    pol = _clean_locode(pol_locode, "POL")
    pod = _clean_locode(pod_locode, "POD")
    if pol is None or pod is None:
        raise HTTPException(status_code=400, detail="POL et POD sont requis pour une route.")
    if pol == pod:
        raise HTTPException(status_code=400, detail="POL et POD doivent être différents.")
    if _match_route(grid, pol, pod) is not None:
        raise HTTPException(status_code=400, detail="Cette route existe déjà sur la grille.")
    manual_base = _opt_decimal(base_rate)
    distance, nav_days, opex_daily, base = await compute_route_economics(
        db,
        pol_locode=pol,
        pod_locode=pod,
        vessel_id=grid.vessel_id,
        distance_nm=_opt_decimal(distance_nm),
    )
    route = RateGridLine(
        grid_id=grid.id,
        pol_locode=pol,
        pod_locode=pod,
        distance_nm=distance,
        nav_days=nav_days,
        opex_daily=opex_daily,
        base_rate=manual_base if manual_base is not None else base,
        is_manual=manual_base is not None,
    )
    db.add(route)
    await db.flush()
    await activity_record(
        db,
        action="create",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid_line",
        entity_id=route.id,
        entity_label=f"{grid.reference} · {pol}→{pod}",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/grids/{grid_id}")


@router.post("/grids/{grid_id}/routes/{route_id}/edit")
async def grid_route_edit(
    grid_id: int,
    route_id: int,
    request: Request,
    distance_nm: str | None = Form(None),
    base_rate: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    """Édite la distance et/ou le base_rate (surcharge manuelle) d'une route."""
    grid = await _grid_editable(db, grid_id)
    route = await db.get(RateGridLine, route_id)
    if route is None or route.grid_id != grid.id:
        raise HTTPException(status_code=404)
    manual_base = _opt_decimal(base_rate)
    distance, nav_days, opex_daily, base = await compute_route_economics(
        db,
        pol_locode=route.pol_locode,
        pod_locode=route.pod_locode,
        vessel_id=grid.vessel_id,
        distance_nm=_opt_decimal(distance_nm),
    )
    route.distance_nm = distance
    route.nav_days = nav_days
    route.opex_daily = opex_daily
    if manual_base is not None:
        route.base_rate = manual_base
        route.is_manual = True
    else:
        route.base_rate = base
        route.is_manual = False
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid_line",
        entity_id=route.id,
        entity_label=f"{grid.reference} · {route.pol_locode}→{route.pod_locode}",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/grids/{grid_id}")


@router.post("/grids/{grid_id}/routes/{route_id}/recalculate")
async def grid_route_recalculate(
    grid_id: int,
    route_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    """Recalcule l'économie OPEX d'une route (efface la surcharge manuelle)."""
    grid = await _grid_editable(db, grid_id)
    route = await db.get(RateGridLine, route_id)
    if route is None or route.grid_id != grid.id:
        raise HTTPException(status_code=404)
    await _recalc_route(db, grid, route)
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid_line",
        entity_id=route.id,
        entity_label=f"{grid.reference} · {route.pol_locode}→{route.pod_locode}",
        detail=f"recalculated base_rate={route.base_rate}",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/grids/{grid_id}")


@router.post("/grids/{grid_id}/routes/{route_id}/delete")
async def grid_route_delete(
    grid_id: int,
    route_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "S")),
):
    grid = await _grid_editable(db, grid_id)
    route = await db.get(RateGridLine, route_id)
    if route is None or route.grid_id != grid.id:
        raise HTTPException(status_code=404)
    label = f"{grid.reference} · {route.pol_locode}→{route.pod_locode}"
    route_pk = route.id
    await db.delete(route)
    await db.flush()
    await activity_record(
        db,
        action="delete",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid_line",
        entity_id=route_pk,
        entity_label=label,
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/grids/{grid_id}")


# ────────────────────────────────────────────── Grid brackets (volume)
@router.post("/grids/{grid_id}/brackets")
async def grid_brackets_update(
    grid_id: int,
    request: Request,
    bracket_max_qty: list[str] = Form(default=[]),
    bracket_coeff: list[str] = Form(default=[]),
    bracket_label: list[str] = Form(default=[]),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    """Remplace les brackets de volume de la grille (coefficients dégressifs)."""
    grid = await _grid_editable(db, grid_id)
    brackets: list[dict] = []
    for max_qty_raw, coeff_raw, label_raw in zip(
        bracket_max_qty, bracket_coeff, bracket_label, strict=False
    ):
        max_qty = _opt_int(max_qty_raw)
        coeff = _opt_decimal(coeff_raw)
        if max_qty is None or coeff is None:
            continue  # ligne vide → ignorée
        if max_qty <= 0 or coeff < 0:
            raise HTTPException(
                status_code=400, detail="bracket invalide : max_qty > 0 et coeff ≥ 0"
            )
        label = (label_raw or "").strip() or f"{max_qty} palettes"
        brackets.append(
            {
                "key": f"q{max_qty}",
                "label": label,
                "max_qty": max_qty,
                "coeff": float(coeff),
            }
        )
    if not brackets:
        raise HTTPException(status_code=400, detail="au moins une bracket est requise")
    brackets.sort(key=lambda b: b["max_qty"])
    grid.brackets_json = json.dumps(brackets)
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid",
        entity_id=grid.id,
        entity_label=grid.reference,
        detail=f"brackets updated ({len(brackets)})",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/grids/{grid_id}")


# ────────────────────────────────────────────── Grid options
@router.post("/grids/{grid_id}/options")
async def grid_option_create(
    grid_id: int,
    request: Request,
    code: str = Form(...),
    label: str = Form(...),
    unit: str = Form(...),
    amount_eur: float = Form(...),
    is_active: bool = Form(False),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    grid = await db.get(RateGrid, grid_id)
    if grid is None:
        raise HTTPException(status_code=404)
    if unit not in RATE_OPTION_UNITS:
        raise HTTPException(status_code=400, detail="unité de tarification invalide")
    code_clean = code.strip().upper().replace(" ", "_")[:40]
    label_clean = label.strip()
    if not code_clean or not label_clean:
        raise HTTPException(status_code=400, detail="code et libellé requis")
    amount = Decimal(str(amount_eur))
    if amount < 0:
        raise HTTPException(status_code=400, detail="montant négatif interdit")
    option = RateGridOption(
        grid_id=grid.id,
        code=code_clean,
        label=label_clean,
        unit=unit,
        amount_eur=amount,
        is_active=is_active,
    )
    db.add(option)
    await db.flush()
    await activity_record(
        db,
        action="create",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid_option",
        entity_id=option.id,
        entity_label=f"{grid.reference} · {option.code}",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/grids/{grid_id}")


@router.post("/grids/{grid_id}/options/{option_id}/toggle")
async def grid_option_toggle(
    grid_id: int,
    option_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    grid = await db.get(RateGrid, grid_id)
    option = await db.get(RateGridOption, option_id)
    if grid is None or option is None or option.grid_id != grid.id:
        raise HTTPException(status_code=404)
    option.is_active = not option.is_active
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid_option",
        entity_id=option.id,
        entity_label=f"{grid.reference} · {option.code}",
        detail="activated" if option.is_active else "deactivated",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/grids/{grid_id}")


@router.post("/grids/{grid_id}/options/{option_id}/delete")
async def grid_option_delete(
    grid_id: int,
    option_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "S")),
):
    grid = await db.get(RateGrid, grid_id)
    option = await db.get(RateGridOption, option_id)
    if grid is None or option is None or option.grid_id != grid.id:
        raise HTTPException(status_code=404)
    option_code = option.code
    option_pk = option.id
    await db.delete(option)
    await db.flush()
    await activity_record(
        db,
        action="delete",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_grid_option",
        entity_id=option_pk,
        entity_label=f"{grid.reference} · {option_code}",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/grids/{grid_id}")


# ────────────────────────────────────────────── Devis (quotes émis)
@router.get("/devis", response_class=HTMLResponse)
async def devis_list(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    """100 derniers devis émis par l'outil public /devis et le booking."""
    rows = (
        await db.execute(
            select(Quote, ClientAccount.email, ClientAccount.company_name)
            .outerjoin(ClientAccount, ClientAccount.id == Quote.client_account_id)
            .order_by(Quote.created_at.desc())
            .limit(100)
        )
    ).all()
    quotes = [
        {"quote": quote, "account_email": email, "account_company": company}
        for quote, email, company in rows
    ]
    return templates.TemplateResponse(
        "staff/commercial/devis_list.html",
        {"request": request, "user": user, "quotes": quotes},
    )


@router.get("/devis/{reference}", response_class=HTMLResponse)
async def devis_detail_staff(
    reference: str,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    """Détail d'un devis côté commercial : lignes, ajustement, consultations."""
    quote = (
        await db.execute(
            select(Quote)
            .where(Quote.reference == reference)
            .options(selectinload(Quote.lines), selectinload(Quote.views))
        )
    ).scalar_one_or_none()
    if quote is None:
        raise HTTPException(status_code=404, detail="Devis introuvable")
    pol = (
        await db.execute(select(Port).where(Port.locode == quote.pol_locode))
    ).scalar_one_or_none()
    pod = (
        await db.execute(select(Port).where(Port.locode == quote.pod_locode))
    ).scalar_one_or_none()
    leg = await db.get(Leg, quote.leg_id) if quote.leg_id else None
    return templates.TemplateResponse(
        "staff/commercial/devis_detail.html",
        {
            "request": request,
            "user": user,
            "quote": quote,
            "pol": pol,
            "pod": pod,
            "leg": leg,
        },
    )


@router.post("/devis/{reference}/adjust")
async def devis_adjust(
    reference: str,
    request: Request,
    adjustment_eur: str = Form("0"),
    adjustment_comment: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    """Applique une remise / majoration (signée) + commentaire sur un devis."""
    quote = (
        await db.execute(select(Quote).where(Quote.reference == reference))
    ).scalar_one_or_none()
    if quote is None:
        raise HTTPException(status_code=404, detail="Devis introuvable")
    raw = (adjustment_eur or "0").strip().replace(",", ".")
    try:
        amount = Decimal(raw)
    except (InvalidOperation, ValueError):
        amount = Decimal("0")
    quote.adjustment_eur = amount
    quote.adjustment_comment = (adjustment_comment or "").strip() or None
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="quote",
        entity_id=quote.id,
        entity_label=quote.reference,
        detail=f"ajustement {amount} EUR",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/devis/{reference}", status_code=303)


# ────────────────────────────────────────────── Offers
@router.get("/offers", response_class=HTMLResponse)
async def offers_list(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    offers = list(
        (await db.execute(select(RateOffer).order_by(RateOffer.created_at.desc()))).scalars().all()
    )
    return templates.TemplateResponse(
        "staff/commercial/offers.html",
        {"request": request, "user": user, "offers": offers},
    )


@router.get("/offers/new", response_class=HTMLResponse)
async def offer_new_form(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
) -> HTMLResponse:
    clients = list(
        (await db.execute(select(Client).where(Client.is_active.is_(True)).order_by(Client.name)))
        .scalars()
        .all()
    )
    from app.services.leg_filter import leg_select_options

    legs = list((await db.execute(select(Leg).order_by(Leg.etd.desc()).limit(50))).scalars().all())
    leg_options = await leg_select_options(db)
    grids = list(
        (
            await db.execute(
                select(RateGrid)
                .options(selectinload(RateGrid.lines))
                .where(RateGrid.status == "active")
            )
        )
        .scalars()
        .all()
    )
    return templates.TemplateResponse(
        "staff/commercial/offer_form.html",
        {
            "request": request,
            "user": user,
            "clients": clients,
            "legs": legs,
            "leg_options": leg_options,
            "grids": grids,
            "offer": None,
        },
    )


async def _grids_for(db: AsyncSession, *, client_id: int | None) -> list[RateGrid]:
    """Grilles actives applicables à un client (multi-routes).

    Retenues : statut ``active``, valides à ce jour, et soit propres au client
    soit grilles par défaut (``client_id`` NULL). La route est résolue à la
    création de l'offre via la ligne-route POL/POD de la grille (cf. le leg
    ciblé). Les grilles spécifiques au client sont listées avant les défauts.
    """
    today = datetime.now(UTC).date()
    query = (
        select(RateGrid)
        .options(selectinload(RateGrid.lines))
        .where(
            RateGrid.status == "active",
            RateGrid.valid_from <= today,
            or_(RateGrid.valid_to.is_(None), RateGrid.valid_to >= today),
        )
    )
    if client_id:
        query = query.where(or_(RateGrid.client_id == client_id, RateGrid.client_id.is_(None)))
    else:
        query = query.where(RateGrid.client_id.is_(None))

    grids = list((await db.execute(query)).scalars().all())
    # Client-specific d'abord, puis défaut ; tri secondaire par référence.
    grids.sort(key=lambda g: (g.client_id is None, g.reference or ""))
    return grids


@router.get("/offers/grid-options", response_class=HTMLResponse)
async def offer_grid_options(
    request: Request,
    client_id: int | None = None,
    leg_id: int | None = None,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    """Partial HTMX : options du <select> grille filtrées par client."""
    grids = await _grids_for(db, client_id=client_id)
    return templates.TemplateResponse(
        "staff/commercial/_grid_options.html",
        {"request": request, "grids": grids},
    )


@router.post("/offers")
async def offer_create(
    request: Request,
    client_id: int = Form(...),
    title: str = Form(...),
    grid_id: int | None = Form(None),
    leg_id: int | None = Form(None),
    estimated_palettes: int = Form(0),
    valid_until: str | None = Form(None),
    notes: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    if not await db.get(Client, client_id):
        raise HTTPException(status_code=404, detail="client introuvable")
    grid = (
        (
            await db.execute(
                select(RateGrid).options(selectinload(RateGrid.lines)).where(RateGrid.id == grid_id)
            )
        ).scalar_one_or_none()
        if grid_id
        else None
    )
    proposed_rate = Decimal("0")
    total = Decimal("0")
    if grid is not None and estimated_palettes > 0 and grid.lines:
        # Route ciblée : ligne POL/POD du leg, sinon première route de la grille.
        route = None
        if leg_id:
            leg = await db.get(Leg, leg_id)
            if leg is not None:
                pol_port = await db.get(Port, leg.departure_port_id)
                pod_port = await db.get(Port, leg.arrival_port_id)
                if pol_port and pod_port:
                    route = _match_route(grid, pol_port.locode, pod_port.locode)
        if route is None:
            route = grid.lines[0]
        # Bracket de volume au niveau grille (coefficients dégressifs).
        picked = pick_bracket(grid.brackets, estimated_palettes)
        if picked:
            proposed_rate = bracket_rate(
                base_rate=route.base_rate,
                coeff=picked["coeff"],
                adjustment_index=grid.adjustment_index,
            )
            total = (proposed_rate * Decimal(estimated_palettes)).quantize(Decimal("0.01"))

    ref = await next_offer_reference(db)
    offer = RateOffer(
        reference=ref,
        client_id=client_id,
        grid_id=grid.id if grid else None,
        leg_id=leg_id,
        title=title.strip(),
        status="draft",
        estimated_palettes=estimated_palettes,
        proposed_rate_eur=proposed_rate or None,
        total_eur=total or None,
        valid_until=_date.fromisoformat(valid_until) if valid_until else None,
        notes=notes,
    )
    db.add(offer)
    await db.flush()
    await activity_record(
        db,
        action="create",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_offer",
        entity_id=offer.id,
        entity_label=ref,
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url="/commercial/offers", status_code=303)


@router.post("/offers/{offer_id}/send")
async def offer_send(
    offer_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    o = await db.get(RateOffer, offer_id)
    if o is None:
        raise HTTPException(status_code=404)
    o.status = "sent"
    o.sent_at = datetime.now(UTC)
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_offer",
        entity_id=o.id,
        entity_label=o.reference,
        detail="sent",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url="/commercial/offers", status_code=303)


@router.post("/offers/{offer_id}/convert")
async def offer_convert_to_order(
    offer_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    """Convertir une offre en commande ferme."""
    offer = await db.get(RateOffer, offer_id)
    if offer is None:
        raise HTTPException(status_code=404)
    if offer.status not in ("draft", "sent", "accepted"):
        raise HTTPException(status_code=400, detail="offre non convertible")
    offer.status = "accepted"
    offer.accepted_at = datetime.now(UTC)
    ref = await next_order_reference(db)
    order = Order(
        reference=ref,
        client_id=offer.client_id,
        offer_id=offer.id,
        leg_id=offer.leg_id,
        status="draft",
        booked_palettes=offer.estimated_palettes or 0,
        rate_per_palette_eur=offer.proposed_rate_eur,
        total_eur=offer.total_eur,
        pipedrive_deal_id=offer.pipedrive_deal_id,
    )
    db.add(order)
    await db.flush()
    await activity_record(
        db,
        action="create",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="order",
        entity_id=order.id,
        entity_label=ref,
        detail=f"from offer {offer.reference}",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/orders/{order.id}", status_code=303)


# ────────────────────────────────────────────── Orders
@router.get("/orders", response_class=HTMLResponse)
async def orders_list(
    request: Request,
    status: str | None = None,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    stmt = select(Order).options(selectinload(Order.client))
    if status:
        stmt = stmt.where(Order.status == status)
    stmt = stmt.order_by(Order.created_at.desc())
    orders = list((await db.execute(stmt)).scalars().all())
    return templates.TemplateResponse(
        "staff/commercial/orders.html",
        {"request": request, "user": user, "orders": orders, "filter_status": status},
    )


@router.get("/orders/new", response_class=HTMLResponse)
async def order_new_form(
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
) -> HTMLResponse:
    clients = list(
        (await db.execute(select(Client).where(Client.is_active.is_(True)).order_by(Client.name)))
        .scalars()
        .all()
    )
    from app.services.leg_filter import leg_select_options

    legs = list((await db.execute(select(Leg).order_by(Leg.etd.desc()).limit(50))).scalars().all())
    leg_options = await leg_select_options(db)
    return templates.TemplateResponse(
        "staff/commercial/order_form.html",
        {
            "request": request,
            "user": user,
            "clients": clients,
            "legs": legs,
            "leg_options": leg_options,
            "order": None,
        },
    )


@router.post("/orders")
async def order_create(
    request: Request,
    client_id: int = Form(...),
    leg_id: int | None = Form(None),
    booked_palettes: int = Form(0),
    rate_per_palette_eur: float | None = Form(None),
    cargo_description: str | None = Form(None),
    shipper_name: str | None = Form(None),
    consignee_name: str | None = Form(None),
    palette_format: str | None = Form(None),
    weight_per_palette_kg: str | None = Form(None),
    thc_included: str | None = Form(None),
    booking_fee: str | None = Form(None),
    documentation_fee: str | None = Form(None),
    departure_locode: str | None = Form(None),
    arrival_locode: str | None = Form(None),
    delivery_date_start: str | None = Form(None),
    delivery_date_end: str | None = Form(None),
    rate_grid_id: str | None = Form(None),
    rate_grid_line_id: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    if not await db.get(Client, client_id):
        raise HTTPException(status_code=404, detail="client introuvable")
    ref = await next_order_reference(db)
    rate = Decimal(str(rate_per_palette_eur)) if rate_per_palette_eur else None
    total = (rate * Decimal(booked_palettes)).quantize(Decimal("0.01")) if rate else None
    order = Order(
        reference=ref,
        client_id=client_id,
        leg_id=leg_id,
        status="draft",
        booked_palettes=booked_palettes,
        rate_per_palette_eur=rate,
        total_eur=total,
        cargo_description=cargo_description,
        shipper_name=shipper_name,
        consignee_name=consignee_name,
        palette_format=(palette_format or "").strip() or None,
        weight_per_palette_kg=_opt_decimal(weight_per_palette_kg),
        thc_included=thc_included in ("on", "true", "1"),
        booking_fee=_opt_decimal(booking_fee),
        documentation_fee=_opt_decimal(documentation_fee),
        departure_locode=_clean_locode(departure_locode, "POL"),
        arrival_locode=_clean_locode(arrival_locode, "POD"),
        delivery_date_start=_opt_date(delivery_date_start, "Livraison (début)"),
        delivery_date_end=_opt_date(delivery_date_end, "Livraison (fin)"),
        rate_grid_id=_opt_int(rate_grid_id),
        rate_grid_line_id=_opt_int(rate_grid_line_id),
    )
    db.add(order)
    await db.flush()
    await activity_record(
        db,
        action="create",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="order",
        entity_id=order.id,
        entity_label=ref,
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/orders/{order.id}", status_code=303)


@router.get("/orders/{order_id}", response_class=HTMLResponse)
async def order_detail(
    order_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> HTMLResponse:
    order = (
        await db.execute(
            select(Order)
            .options(selectinload(Order.client), selectinload(Order.assignments))
            .where(Order.id == order_id)
        )
    ).scalar_one_or_none()
    if order is None:
        raise HTTPException(status_code=404)
    leg = await db.get(Leg, order.leg_id) if order.leg_id else None
    return templates.TemplateResponse(
        "staff/commercial/order_detail.html",
        {"request": request, "user": user, "order": order, "leg": leg},
    )


@router.get("/orders/{order_id}/assign", response_class=HTMLResponse)
async def order_assign_form(
    order_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
) -> HTMLResponse:
    """COM-01 — écran d'affectation : legs compatibles avec la route souhaitée,
    suggestion automatique et alerte « hors délai » (ETA > date de livraison).

    Affectation **simple-leg** (parité V2) : une commande est affectée à un
    seul leg ; réaffecter remplace l'affectation. La ventilation multi-legs
    avec répartition du CA est un chantier P1 distinct (cf. backlog COM).
    """
    order = (
        await db.execute(
            select(Order)
            .options(selectinload(Order.client), selectinload(Order.assignments))
            .where(Order.id == order_id)
        )
    ).scalar_one_or_none()
    if order is None:
        raise HTTPException(status_code=404)

    legs = await compatible_legs_for_order(db, order)
    current_leg_id = order.assignments[0].leg_id if order.assignments else None
    # Pré-sélection : le leg actuellement affecté s'il figure dans les
    # candidats, sinon la suggestion automatique (1er compatible dans les délais).
    suggested = next((lg for lg in legs if lg.id == current_leg_id), None)
    if suggested is None:
        suggested = suggest_leg_for_order(legs, order)
    ports = {p.id: p for p in (await db.execute(select(Port))).scalars().all()}
    leg_rows = [
        {
            "leg": lg,
            "pol": ports.get(lg.departure_port_id),
            "pod": ports.get(lg.arrival_port_id),
            "late": leg_is_late_for_order(lg, order),
            "current": lg.id == current_leg_id,
        }
        for lg in legs
    ]
    return templates.TemplateResponse(
        "staff/commercial/assign_form.html",
        {
            "request": request,
            "user": user,
            "order": order,
            "leg_rows": leg_rows,
            "suggested": suggested,
        },
    )


@router.post("/orders/{order_id}/assign")
async def order_assign_submit(
    order_id: int,
    request: Request,
    leg_id: int = Form(...),
    notes: str | None = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    """Affecte la commande à un leg (simple-leg, parité V2). Réaffecter
    remplace l'affectation existante. Les palettes affectées dérivent de la
    commande (``booked_palettes``) pour rester cohérentes avec la capacité —
    pas de quantité divergente saisie ici.
    """
    order = (
        await db.execute(
            select(Order)
            .options(selectinload(Order.assignments))
            .where(Order.id == order_id)
        )
    ).scalar_one_or_none()
    if order is None:
        raise HTTPException(status_code=404)
    leg = await db.get(Leg, leg_id)
    if leg is None:
        raise HTTPException(status_code=404, detail="leg introuvable")
    if leg.atd is not None:
        raise HTTPException(
            status_code=400, detail="ce leg est déjà parti — affectation impossible"
        )

    # Simple-leg : on remplace toute affectation existante (sur ce leg ou un autre).
    for existing in list(order.assignments):
        await db.delete(existing)
    await db.flush()
    db.add(
        OrderAssignment(
            order_id=order.id,
            leg_id=leg_id,
            palettes_count=max(0, order.booked_palettes),
            pallet_format=(order.palette_format or "EPAL").strip() or "EPAL",
            notes=(notes or "").strip() or None,
        )
    )
    order.leg_id = leg_id  # le reste de l'app (packing list, stowage) lit order.leg_id
    # NB — l'affectation est orthogonale à la confirmation commerciale : on ne
    # touche pas au statut ici (la confirmation passe par ``order_confirm`` qui
    # pose ``confirmed_at`` et déclenche ses effets de bord).
    await db.flush()
    await activity_record(
        db,
        action="assign",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="order",
        entity_id=order.id,
        entity_label=order.reference,
        detail=f"→ leg {leg.leg_code} ({order.booked_palettes} pal.)",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/orders/{order.id}")


@router.post("/orders/{order_id}/assignments/{assignment_id}/delete")
async def order_assignment_delete(
    order_id: int,
    assignment_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    assignment = await db.get(OrderAssignment, assignment_id)
    if assignment is None or assignment.order_id != order_id:
        raise HTTPException(status_code=404)
    leg_id = assignment.leg_id
    await db.delete(assignment)
    await db.flush()
    # Si le leg principal était celui-ci, le ré-aligner sur une affectation
    # restante (ou le détacher).
    order = (
        await db.execute(
            select(Order)
            .options(selectinload(Order.assignments))
            .where(Order.id == order_id)
        )
    ).scalar_one_or_none()
    if order is not None and order.leg_id == leg_id:
        order.leg_id = order.assignments[0].leg_id if order.assignments else None
        await db.flush()
    await activity_record(
        db,
        action="unassign",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="order",
        entity_id=order_id,
        entity_label=order.reference if order else str(order_id),
        detail=f"retrait affectation leg {leg_id}",
        ip_address=_client_ip(request),
    )
    return _hx_or_redirect(request, f"/commercial/orders/{order_id}")


@router.post("/orders/{order_id}/confirm")
async def order_confirm(
    order_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "M")),
):
    order = await db.get(Order, order_id)
    if order is None:
        raise HTTPException(status_code=404)
    order.status = "confirmed"
    order.confirmed_at = datetime.now(UTC)
    await db.flush()
    await activity_record(
        db,
        action="update",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="order",
        entity_id=order.id,
        entity_label=order.reference,
        detail="confirmed",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url=f"/commercial/orders/{order_id}", status_code=303)


@router.post("/orders/{order_id}/cancel")
async def order_cancel(
    order_id: int,
    request: Request,
    reason: str = Form(""),
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "S")),
):
    order = await db.get(Order, order_id)
    if order is None:
        raise HTTPException(status_code=404)
    order.status = "cancelled"
    order.cancelled_at = datetime.now(UTC)
    order.cancelled_reason = reason or None
    await db.flush()
    await activity_record(
        db,
        action="delete",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="order",
        entity_id=order.id,
        entity_label=order.reference,
        detail=f"cancelled: {reason}",
        ip_address=_client_ip(request),
    )
    return RedirectResponse(url="/commercial/orders", status_code=303)


# ────────────────────────────────────────────── Offer DOCX export
@router.get("/offers/{offer_id}/export.docx")
async def offer_export_docx(
    offer_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(require_permission("commercial", "C")),
) -> Response:
    """Exporte une offre commerciale en fichier Word (.docx)."""
    offer = await db.get(RateOffer, offer_id)
    if offer is None:
        raise HTTPException(status_code=404, detail="offre introuvable")

    client = await db.get(Client, offer.client_id)
    leg = await db.get(Leg, offer.leg_id) if offer.leg_id else None

    # ── Build document ──────────────────────────────────────────────────────
    doc = Document()

    # ── Header: OFFRE COMMERCIALE + reference
    heading = doc.add_heading("", level=0)
    run = heading.add_run("OFFRE COMMERCIALE NEWTOWT")
    run.font.color.rgb = RGBColor(0x0D, 0x59, 0x66)
    run.font.size = Pt(20)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_run = ref_para.add_run(f"Référence : {offer.reference}")
    ref_run.bold = True
    ref_run.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # ── Section Client ──────────────────────────────────────────────────────
    client_heading = doc.add_heading("Client", level=2)
    client_heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x59, 0x66)

    client_table = doc.add_table(rows=0, cols=2)
    client_table.style = "Table Grid"

    def _add_kv_row(table, label: str, value: str) -> None:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value

    _add_kv_row(client_table, "Nom", client.name if client else "—")
    if client and client.company_name:
        _add_kv_row(client_table, "Société", client.company_name)
    _add_kv_row(client_table, "E-mail", client.email if client else "—")
    _add_kv_row(client_table, "Téléphone", client.phone if client else "—")

    doc.add_paragraph()  # spacer

    # ── Section Objet ───────────────────────────────────────────────────────
    objet_heading = doc.add_heading("Objet", level=2)
    objet_heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x59, 0x66)
    doc.add_paragraph(offer.title or "—")

    doc.add_paragraph()  # spacer

    # ── Section Itinéraire ──────────────────────────────────────────────────
    itin_heading = doc.add_heading("Itinéraire", level=2)
    itin_heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x59, 0x66)

    if leg:
        etd_str = leg.etd.strftime("%d/%m/%Y") if leg.etd else "—"
        eta_str = leg.eta.strftime("%d/%m/%Y") if leg.eta else "—"
        doc.add_paragraph(f"Leg : {leg.leg_code}\n" f"ETD : {etd_str}     ETA : {eta_str}")
    else:
        doc.add_paragraph("À confirmer")

    doc.add_paragraph()  # spacer

    # ── Table Tarification ──────────────────────────────────────────────────
    tarif_heading = doc.add_heading("Tarification", level=2)
    tarif_heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x59, 0x66)

    tarif_table = doc.add_table(rows=1, cols=4)
    tarif_table.style = "Table Grid"
    hdr_cells = tarif_table.rows[0].cells
    for idx, label in enumerate(["Description", "Quantité", "Tarif unitaire", "Total"]):
        hdr_cells[idx].text = label
        hdr_cells[idx].paragraphs[0].runs[0].bold = True

    row_cells = tarif_table.add_row().cells
    qty = offer.estimated_palettes or 0
    rate = offer.proposed_rate_eur
    total = offer.total_eur
    rate_str = f"{rate:,.2f} EUR/palette".replace(",", " ") if rate is not None else "—"
    total_str = f"{total:,.2f} EUR".replace(",", " ") if total is not None else "—"
    row_cells[0].text = "Fret aérien palettes"
    row_cells[1].text = str(qty)
    row_cells[2].text = rate_str
    row_cells[3].text = total_str

    doc.add_paragraph()  # spacer

    # ── Section Conditions ──────────────────────────────────────────────────
    cond_heading = doc.add_heading("Conditions", level=2)
    cond_heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x59, 0x66)

    validity_str = offer.valid_until.strftime("%d/%m/%Y") if offer.valid_until else "—"
    cond_para = doc.add_paragraph()
    cond_para.add_run(f"Validité : {validity_str}\n").bold = False
    cond_para.add_run(
        "Ce prix inclut le transport par voilier cargo à propulsion vélique "
        "(zéro émission directe)."
    )

    # ── Section Notes (optionnelle) ─────────────────────────────────────────
    if offer.notes:
        doc.add_paragraph()  # spacer
        notes_heading = doc.add_heading("Notes", level=2)
        notes_heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x59, 0x66)
        doc.add_paragraph(offer.notes)

    doc.add_paragraph()  # spacer

    # ── Footer ──────────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph(
        "NEWTOWT — Pioneer of wind-powered cargo since 2011 — www.newtowt.eu"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x59, 0x66)
        run.font.size = Pt(9)
        run.italic = True

    # ── Serialize ───────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    await activity_record(
        db,
        action="offer_export_docx",
        user_id=user.id,
        user_name=user.full_name or user.username,
        user_role=user.role,
        module="commercial",
        entity_type="rate_offer",
        entity_id=offer.id,
        entity_label=offer.reference,
        detail=offer.reference,
        ip_address=_client_ip(request),
    )

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="Offre_{offer.reference}.docx"'},
    )


def _client_ip(request: Request) -> str | None:
    return request.headers.get("x-forwarded-for") or (
        request.client.host if request.client else None
    )
